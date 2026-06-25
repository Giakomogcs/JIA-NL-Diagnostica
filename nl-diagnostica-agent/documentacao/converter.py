#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Conversor Markdown -> DOCX com identidade visual NL Diagnostica (laranja institucional).
Gera documentacao/NL-Diagnostica-Documentacao.docx a partir de DOCUMENTACAO.md.
"""
import os
import re
import shutil
import subprocess
import tempfile
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SRC = os.path.join(HERE, "DOCUMENTACAO.md")
OUT = os.path.join(HERE, "NL-Diagnostica-Documentacao.docx")

# Paleta NL Diagnostica
ACCENT = "F7941E"   # laranja institucional
ACCENT2 = "B85C0A"  # laranja escuro
INK = "1F2A37"      # titulos escuros
INK2 = "1F2937"     # corpo
GRAY = "6B7785"     # legendas
RULE = "D9DEE6"     # bordas
ZEBRA = "FFF7EC"    # linhas alternadas (tint laranja)
CELLHDR = "FFE7C7"  # cabecalho de tabela
CODE_BG = "0B0B0B"
CODE_FG = "E5E7EB"
CALLOUT = {
    "NOTE": ("1D4ED8", "EAF0FE", "i"),
    "TIP": ("047857", "E7F6F0", "+"),
    "WARNING": ("B45309", "FEF3E2", "!"),
    "DANGER": ("B91C1C", "FDECEC", "x"),
}
FONT = "Calibri"
MONO = "Consolas"

MMDC = shutil.which("mmdc")


def set_cell_bg(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_border(cell, color=RULE, sz=4, sides=("top", "bottom", "left", "right")):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in sides:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom), ("start", left), ("end", right)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        m.append(el)
    tcPr.append(m)


def shade_paragraph(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def para_border(p, color, size=24, side="left"):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "8")
    el.set(qn("w:color"), color)
    pbdr.append(el)
    pPr.append(pbdr)


def bottom_rule(p, color, size=8):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    el = OxmlElement("w:bottom")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), str(size))
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), color)
    pbdr.append(el)
    pPr.append(pbdr)


INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def add_runs(p, text, base_color=INK2, base_size=10.5, base_bold=False):
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1]); r.font.name = MONO
            r.font.color.rgb = RGBColor.from_string(ACCENT2)
            r.font.size = Pt(base_size - 0.5)
        elif part.startswith("*") and part.endswith("*"):
            r = p.add_run(part[1:-1]); r.italic = True
        else:
            r = p.add_run(part)
        r.font.name = part and r.font.name or FONT
        if r.font.name is None:
            r.font.name = FONT
        if r.font.color.rgb is None:
            r.font.color.rgb = RGBColor.from_string(base_color)
        if r.font.size is None:
            r.font.size = Pt(base_size)
        if base_bold:
            r.bold = True


def add_field(p, instr):
    r = p.add_run()
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin"); r._r.append(fb)
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr; r._r.append(it)
    fs = OxmlElement("w:fldChar"); fs.set(qn("w:fldCharType"), "separate"); r._r.append(fs)
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end"); r._r.append(fe)


def render_mermaid(code, idx):
    if not MMDC:
        return None
    try:
        tmp = tempfile.mkdtemp()
        mmd = os.path.join(tmp, f"d{idx}.mmd")
        png = os.path.join(tmp, f"d{idx}.png")
        with open(mmd, "w", encoding="utf-8") as f:
            f.write(code)
        subprocess.run([MMDC, "-i", mmd, "-o", png, "-b", "white", "-s", "2"],
                       check=True, capture_output=True, timeout=120)
        return png if os.path.exists(png) else None
    except Exception:
        return None


# ----- cover, header, footer -----

def build_cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_rule(p, ACCENT, size=36)
    r = p.add_run("NL DIAGNOSTICA AGENT")
    r.font.name = FONT; r.font.size = Pt(34); r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p2.add_run("Documentacao Tecnica e de Negocio")
    r.font.name = FONT; r.font.size = Pt(16); r.font.color.rgb = RGBColor.from_string(INK)
    p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p3.add_run("Triagem inteligente de editais de licitacao  |  n8n + Supabase + Azure OpenAI")
    r.font.name = FONT; r.font.size = Pt(11); r.font.color.rgb = RGBColor.from_string(GRAY)
    for _ in range(6):
        doc.add_paragraph()
    meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_rule(meta, ACCENT2, size=12)
    r = meta.add_run("Gerado por Doc Master")
    r.font.name = FONT; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT2)
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = d.add_run("Junho de 2026")
    r.font.name = FONT; r.font.size = Pt(10); r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def build_toc(doc):
    p = doc.add_paragraph()
    r = p.add_run("Sumario")
    r.font.name = FONT; r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    bottom_rule(p, RULE)
    tp = doc.add_paragraph()
    add_field(tp, 'TOC \\o "1-3" \\h \\z \\u')
    note = doc.add_paragraph()
    r = note.add_run("(Clique com o botao direito e escolha 'Atualizar campo' para preencher o sumario.)")
    r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
    doc.add_page_break()


def add_page_field(p, instr):
    add_field(p, instr)


def build_header_footer(section, title, project):
    header = section.header
    hp = header.paragraphs[0]
    hp.text = ""
    tab = hp.paragraph_format
    r = hp.add_run(title)
    r.font.name = FONT; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
    r2 = hp.add_run("\t\t" + project)
    r2.font.name = FONT; r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor.from_string(ACCENT2); r2.bold = True
    bottom_rule(hp, RULE)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = ""
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = fp.add_run("Pagina ")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(fp, "PAGE")
    r = fp.add_run(" de ")
    r.font.size = Pt(8.5); r.font.color.rgb = RGBColor.from_string(GRAY)
    add_page_field(fp, "NUMPAGES")


# ----- block renderers -----

def add_heading(doc, level, text):
    p = doc.add_paragraph()
    if level == 1:
        r = p.add_run(text.upper())
        r.font.name = FONT; r.font.size = Pt(18); r.bold = True
        r.font.color.rgb = RGBColor.from_string(ACCENT)
        p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
        bottom_rule(p, ACCENT, size=12)
    elif level == 2:
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(14); r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK)
        p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
        bottom_rule(p, RULE, size=6)
    else:
        r = p.add_run(text)
        r.font.name = FONT; r.font.size = Pt(12); r.bold = True
        r.font.color.rgb = RGBColor.from_string(INK2)
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    # heading style for TOC
    p.style = doc.styles[f"Heading {min(level,3)}"]
    # re-apply run colors (style may override)
    for r in p.runs:
        r.font.name = FONT
        r.font.color.rgb = RGBColor.from_string(ACCENT if level == 1 else (INK if level == 2 else INK2))
    return p


def add_table(doc, rows):
    header = rows[0]
    body = rows[1:]
    ncol = len(header)
    t = doc.add_table(rows=1, cols=ncol)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    hdr = t.rows[0].cells
    for i, htext in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        add_runs(p, htext, base_color="FFFFFF", base_size=10, base_bold=True)
        for r in p.runs:
            r.font.color.rgb = RGBColor.from_string("FFFFFF"); r.bold = True
        set_cell_bg(hdr[i], ACCENT)
        set_cell_border(hdr[i])
        set_cell_margins(hdr[i])
    for ri, row in enumerate(body):
        cells = t.add_row().cells
        for ci in range(ncol):
            val = row[ci] if ci < len(row) else ""
            cells[ci].text = ""
            p = cells[ci].paragraphs[0]
            add_runs(p, val, base_size=9.5)
            if ri % 2 == 1:
                set_cell_bg(cells[ci], ZEBRA)
            set_cell_border(cells[ci])
            set_cell_margins(cells[ci])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, code, caption=None):
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, CODE_BG)
    set_cell_border(cell, color=CODE_BG)
    set_cell_margins(cell, top=120, bottom=120, left=140, right=140)
    cell.text = ""
    first = True
    for line in code.split("\n"):
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
        r = p.add_run(line if line else " ")
        r.font.name = MONO; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(CODE_FG)
    if caption:
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(caption); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(GRAY)


def add_mermaid(doc, code, idx):
    png = render_mermaid(code, idx)
    if png:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(png, width=Inches(6.2))
        except Exception:
            add_code_block(doc, code, "(diagrama Mermaid)")
            return
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = c.add_run(f"Figura {idx} - diagrama"); r.italic = True; r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor.from_string(GRAY)
    else:
        add_code_block(doc, code, "(diagrama Mermaid - renderize com mermaid-cli para visualizar)")


def add_callout(doc, kind, lines):
    color, bg, icon = CALLOUT.get(kind, CALLOUT["NOTE"])
    t = doc.add_table(rows=1, cols=1)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_border(cell, color=bg)
    set_cell_margins(cell, top=80, bottom=80, left=140, right=140)
    cell.text = ""
    title = cell.paragraphs[0]
    para_border(title, color, size=24, side="left")
    r = title.add_run(f"{icon}  {kind}")
    r.bold = True; r.font.name = FONT; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string(color)
    for ln in lines:
        p = cell.add_paragraph()
        para_border(p, color, size=24, side="left")
        add_runs(p, ln, base_size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_bullet(doc, text, ordered=False, num=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    if ordered:
        r = p.add_run(f"{num}. ")
        r.bold = True; r.font.color.rgb = RGBColor.from_string(ACCENT2); r.font.name = FONT; r.font.size = Pt(10.5)
    else:
        r = p.add_run("\u2022  ")
        r.font.color.rgb = RGBColor.from_string(ACCENT); r.font.name = FONT; r.font.size = Pt(10.5)
    add_runs(p, text, base_size=10.5)


def add_paragraph_md(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text, base_size=10.5)


# ----- markdown parser -----

def parse_table_block(lines):
    rows = []
    for ln in lines:
        ln = ln.strip()
        if set(ln.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            continue  # separator row
        cells = [c.strip() for c in ln.strip("|").split("|")]
        rows.append(cells)
    return rows


def main():
    with open(SRC, encoding="utf-8") as f:
        text = f.read()
    lines = text.split("\n")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(10.5)
    style.font.color.rgb = RGBColor.from_string(INK2)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.9); sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(0.9); sec.right_margin = Inches(0.9)

    build_cover(doc)
    build_toc(doc)
    build_header_footer(sec, "NL Diagnostica Agent - Documentacao", "NL Diagnostica")

    i = 0
    mermaid_idx = 0
    n = len(lines)
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped == "---":
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            if title == "NL Diagnóstica Agent — Documentação Técnica e de Negócio":
                i += 1
                continue
            add_heading(doc, min(level, 3), title)
            i += 1
            continue

        # code / mermaid fences
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            j = i + 1
            buf = []
            while j < n and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            code = "\n".join(buf)
            if lang == "mermaid":
                mermaid_idx += 1
                add_mermaid(doc, code, mermaid_idx)
            else:
                add_code_block(doc, code)
            i = j + 1
            continue

        # callouts
        cm = re.match(r"^>\s*\[!(\w+)\]\s*$", stripped)
        if cm:
            kind = cm.group(1).upper()
            j = i + 1
            buf = []
            while j < n and lines[j].strip().startswith(">"):
                content = lines[j].strip()[1:].strip()
                if content:
                    buf.append(content)
                j += 1
            add_callout(doc, kind, buf)
            i = j
            continue

        # blockquote (plain)
        if stripped.startswith(">"):
            j = i
            buf = []
            while j < n and lines[j].strip().startswith(">"):
                buf.append(lines[j].strip()[1:].strip())
                j += 1
            add_callout(doc, "NOTE", [" ".join(buf)])
            i = j
            continue

        # tables
        if stripped.startswith("|"):
            j = i
            tbl = []
            while j < n and lines[j].strip().startswith("|"):
                tbl.append(lines[j])
                j += 1
            rows = parse_table_block(tbl)
            if rows:
                add_table(doc, rows)
            i = j
            continue

        # ordered list
        om = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if om:
            add_bullet(doc, om.group(2), ordered=True, num=om.group(1))
            i += 1
            continue

        # unordered list
        if stripped.startswith("- ") or stripped.startswith("* "):
            add_bullet(doc, stripped[2:])
            i += 1
            continue

        # plain paragraph
        add_paragraph_md(doc, stripped)
        i += 1

    doc.save(OUT)
    print(f"OK: {OUT}")
    print(f"mermaid renderer: {'mmdc' if MMDC else 'nao encontrado (diagramas como bloco de codigo)'}")


if __name__ == "__main__":
    main()
